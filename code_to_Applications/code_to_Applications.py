
# -*- coding: cp1251 -*-

from docx import Document
from encodings import utf_8
import docx
from docx.shared import Pt
from docx.shared import Mm
from docx.enum.style import WD_STYLE_TYPE
import os

print('Enter extension files (ex: ".m", ".py" etc.): ')
ext = input()

listfiles = sorted(os.listdir())

doc = docx.Document()

 # Стиль заголовка
style3 = doc.styles['Header']
style3.font.name = 'Times New Roman'
style3.font.size = Pt(14)
style3.font.bold = True

 # Стиль обычных абзацев
style1 = doc.styles['Normal']
style1.font.name = 'Times New Roman'
style1.font.size = Pt(14)

 # Стиль кода
style2 = doc.styles.add_style('Codee', WD_STYLE_TYPE.PARAGRAPH)
style2.font.name = 'Courier New'
style2.font.size = Pt(14)

count = 1;
for file in listfiles:

    if file[-len(ext):len(file)] == ext:

        par1 = doc.add_paragraph('ПРИЛОЖЕНИЕ ' + str(count))
        par1.style = doc.styles['Header']
        p_fmt = par1.paragraph_format
        p_fmt.first_line_indent = Mm(12.5)
        p_fmt.line_spacing = 1.5
        p_fmt.space_after = Pt(10)
        par2 = doc.add_paragraph(file)
        par2.style = doc.styles['Normal']
    
        p_fmt3 = par2.paragraph_format
        p_fmt3.first_line_indent = Mm(12.5)
        p_fmt3.line_spacing = 1.5

        try:  
            f = open('' + file,'r', encoding='utf-8')
            print(f)
            for st in f:
                if len(st) > 5:
                    st = st.replace('\n','')
                    cod1 = doc.add_paragraph(st)

                    cod1.style = doc.styles['Codee']
                    p_fmt1 = cod1.paragraph_format
                    p_fmt3.left_indent = Mm(0)
                    p_fmt3.first_line_indent = Mm(12.5)
                    p_fmt1.line_spacing = 1
                    p_fmt1.space_before = Pt(0)
                    p_fmt1.space_after = Pt(0)

            f.close()
    
        
        except:
            print(file + ' : NaN File')

        doc.add_page_break()
        count += 1
if count > 1:
    try:
        doc.save('app.docx')
    except:
        print('ERROR! There is no access to the file.')
else:
    print('ERROR! No files found!')

